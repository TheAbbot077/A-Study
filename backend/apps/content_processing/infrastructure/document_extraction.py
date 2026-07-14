from __future__ import annotations

import hashlib
import io
import os
import re
import shutil
import zipfile

from django.conf import settings

from apps.content_processing.application.extraction_ports import BlockEvidence, ExtractionEvidence, InspectionEvidence, SourceDocument
from apps.content_processing.domain.extraction import (
    DocumentTextClassification, EvidenceOrigin, ExtractedBlockType, ExtractionMethod, NativeTextQuality, OcrRequirement, SourceFormat, sanitize_source_text,
)
from apps.storage.domain.models import StoredFile
from apps.storage.infrastructure.providers import LocalStorageProvider


INSPECTOR_VERSION = "6c2-inspector-1"
EXTRACTOR_VERSION = "6c2-extractor-1"


class ExtractionPolicy:
    minimum_page_characters = 20
    high_quality_page_characters = 200
    maximum_ocr_pages = 100
    maximum_blocks = 50_000
    maximum_characters = 20_000_000


def detect_source_format(source: SourceDocument) -> tuple[str, str, list[dict[str, object]]]:
    extension = os.path.splitext(source.filename)[1].lower()
    declared = (source.declared_content_type or "").lower()
    warnings: list[dict[str, object]] = []
    if source.content.startswith(b"%PDF-"):
        detected, mime = SourceFormat.PDF, "application/pdf"
    elif source.content.startswith(b"PK") and zipfile.is_zipfile(io.BytesIO(source.content)):
        with zipfile.ZipFile(io.BytesIO(source.content)) as package:
            names = set(package.namelist())
        detected = SourceFormat.DOCX if "[Content_Types].xml" in names and "word/document.xml" in names else SourceFormat.UNKNOWN
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if detected == SourceFormat.DOCX else "application/zip"
    else:
        detected, mime = SourceFormat.UNKNOWN, "application/octet-stream"
    expected_extension = {SourceFormat.PDF: ".pdf", SourceFormat.DOCX: ".docx"}.get(detected)
    if expected_extension and extension != expected_extension:
        warnings.append({"code": "extension_signature_mismatch", "message": "Filename extension does not match the detected document signature."})
    if declared and mime != declared:
        warnings.append({"code": "mime_signature_mismatch", "message": "Declared MIME type does not match the detected document signature."})
    return detected, mime, warnings


class StoragePlatformSourceReader:
    def __init__(self, provider=None) -> None:
        self.provider = provider or LocalStorageProvider(settings.MEDIA_ROOT)

    def read(self, stored_file_id: str) -> SourceDocument:
        stored = StoredFile.objects.get(id=stored_file_id)
        stream = self.provider.download(stored.stored_filename)
        try:
            stream.seek(0)
            content = stream.read()
        finally:
            stream.close()
        checksum = stored.checksum or hashlib.sha256(content).hexdigest()
        return SourceDocument(str(stored.id), stored.original_filename, stored.content_type or "", stored.size_bytes, checksum, content)


class PdfDocumentInspector:
    name, version = "pypdf-inspector", INSPECTOR_VERSION

    def supports(self, detected_format) -> bool:
        return detected_format == SourceFormat.PDF

    def inspect(self, source: SourceDocument) -> InspectionEvidence:
        from pypdf import PdfReader
        try:
            reader = PdfReader(io.BytesIO(source.content), strict=False)
            encrypted = bool(reader.is_encrypted)
            if encrypted and reader.decrypt("") == 0:
                return InspectionEvidence(SourceFormat.PDF, "application/pdf", None, True, True, False, confidence=.99)
            counts = [len(re.sub(r"\s+", "", page.extract_text() or "")) for page in reader.pages]
        except Exception:
            return InspectionEvidence(SourceFormat.PDF, "application/pdf", None, corrupt=True, confidence=.95)
        usable = [index for index, count in enumerate(counts) if count >= ExtractionPolicy.minimum_page_characters]
        weak = [index for index, count in enumerate(counts) if count < ExtractionPolicy.minimum_page_characters]
        if counts and len(usable) == len(counts):
            classification, requirement = DocumentTextClassification.NATIVE_TEXT, OcrRequirement.NOT_REQUIRED
        elif usable:
            classification, requirement = DocumentTextClassification.MIXED, OcrRequirement.RECOMMENDED
        elif counts:
            classification, requirement = DocumentTextClassification.SCANNED, OcrRequirement.REQUIRED
        else:
            classification, requirement = DocumentTextClassification.EMPTY_OR_UNRESOLVED, OcrRequirement.REQUIRED
        average = sum(counts) / len(counts) if counts else 0
        quality = NativeTextQuality.HIGH if average >= ExtractionPolicy.high_quality_page_characters else NativeTextQuality.MODERATE if usable else NativeTextQuality.NONE
        return InspectionEvidence(SourceFormat.PDF, "application/pdf", len(counts), native_text_available=bool(usable), native_text_quality=quality, text_classification=classification, ocr_requirement=requirement, ocr_pages_recommended=tuple(weak), parser_recommendation="pypdf-native-with-selective-ocr" if weak else "pypdf-native", confidence=.95)


class DocxDocumentInspector:
    name, version = "python-docx-inspector", INSPECTOR_VERSION

    def supports(self, detected_format) -> bool:
        return detected_format == SourceFormat.DOCX

    def inspect(self, source: SourceDocument) -> InspectionEvidence:
        from docx import Document
        try:
            document = Document(io.BytesIO(source.content))
            text = "".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            return InspectionEvidence(SourceFormat.DOCX, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", None, corrupt=True, confidence=.95)
        available = bool(re.search(r"\w", text)) or bool(document.tables)
        return InspectionEvidence(SourceFormat.DOCX, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", None, native_text_available=available, native_text_quality=NativeTextQuality.HIGH if available else NativeTextQuality.NONE, text_classification=DocumentTextClassification.NATIVE_TEXT if available else DocumentTextClassification.EMPTY_OR_UNRESOLVED, ocr_requirement=OcrRequirement.NOT_REQUIRED, parser_recommendation="python-docx", confidence=.98)


def _page_reference(index: int, page) -> dict[str, object]:
    box = page.mediabox
    return {"page_index": index, "page_number": index + 1, "page_width": float(box.width), "page_height": float(box.height), "rotation": int(page.rotation or 0)}


class PdfDocumentExtractor:
    name, version = "pypdf-layout-extractor", EXTRACTOR_VERSION

    def supports(self, detected_format) -> bool:
        return detected_format == SourceFormat.PDF

    def extract(self, source: SourceDocument, profile) -> ExtractionEvidence:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(source.content), strict=False)
        blocks: list[BlockEvidence] = []
        native_pages = 0
        for page_index, page in enumerate(reader.pages):
            spans: list[dict[str, object]] = []
            def visit(text, cm, tm, font, size):
                clean = sanitize_source_text(text).strip()
                if clean:
                    spans.append({"text": clean, "x": float(tm[4]), "y": float(tm[5]), "font": (font or {}).get("/BaseFont", ""), "size": float(size or 0)})
            page.extract_text(visitor_text=visit)
            lines: dict[int, list[dict[str, object]]] = {}
            for span in spans:
                lines.setdefault(round(float(span["y"]) / 3), []).append(span)
            page_blocks = []
            for line in sorted(lines.values(), key=lambda value: (-float(value[0]["y"]), float(value[0]["x"]))):
                ordered = sorted(line, key=lambda value: float(value["x"]))
                text = " ".join(str(value["text"]) for value in ordered).strip()
                if text:
                    page_blocks.append((text, ordered))
            if page_blocks:
                native_pages += 1
            for page_sequence, (text, line_spans) in enumerate(page_blocks):
                block_type = ExtractedBlockType.PAGE_NUMBER if re.fullmatch(r"\s*\d+\s*", text) else ExtractedBlockType.PARAGRAPH
                x0 = min(float(span["x"]) for span in line_spans); y0 = min(float(span["y"]) for span in line_spans)
                font_size = max(float(span["size"]) for span in line_spans); x1 = max(float(span["x"]) + len(str(span["text"])) * float(span["size"]) * .5 for span in line_spans)
                font_name = str(line_spans[0]["font"])
                geometry = {"x0": x0, "y0": y0, "x1": x1, "y1": y0 + font_size, "coordinate_space": "pdf_points"}
                typography = {"font_family": font_name, "font_size": font_size, "bold": "bold" in font_name.lower(), "italic": "italic" in font_name.lower() or "oblique" in font_name.lower()}
                blocks.append(BlockEvidence(len(blocks), page_sequence, block_type, EvidenceOrigin.LAYOUT_INFERRED, text, re.sub(r"[ \t]+", " ", text).strip(), _page_reference(page_index, page), geometry, typography, source_method="pypdf_visitor", confidence=.88))
            for image_index, image in enumerate(getattr(page, "images", ())):
                blocks.append(BlockEvidence(len(blocks), len(page_blocks) + image_index, ExtractedBlockType.IMAGE, EvidenceOrigin.SOURCE_EXPLICIT, page_reference=_page_reference(page_index, page), source_method="pypdf_image", image_reference=f"page-{page_index + 1}-image-{image_index + 1}", confidence=.9))
        ocr_page_indexes = list(profile.ocr_pages_recommended or [])
        ocr_count = 0
        ocr_engine = ""
        warnings: list[dict[str, object]] = []
        if ocr_page_indexes:
            if not shutil.which("tesseract") or not (shutil.which("pdftoppm") or shutil.which("pdftocairo")):
                if profile.ocr_requirement == OcrRequirement.REQUIRED:
                    raise RuntimeError("OCR unavailable: Tesseract or the PDF rasterizer is missing.")
                warnings.append({"code": "ocr_unavailable", "message": "OCR was recommended but is unavailable; native evidence was preserved."})
            else:
                from pdf2image import convert_from_bytes
                import pytesseract
                for page_index in ocr_page_indexes[:ExtractionPolicy.maximum_ocr_pages]:
                    images = convert_from_bytes(source.content, first_page=page_index + 1, last_page=page_index + 1)
                    if not images:
                        continue
                    data = pytesseract.image_to_data(images[0], output_type=pytesseract.Output.DICT)
                    groups: dict[tuple[int, int, int], list[int]] = {}
                    for index, text in enumerate(data.get("text", [])):
                        if text and text.strip():
                            key = (data["block_num"][index], data["par_num"][index], data["line_num"][index])
                            groups.setdefault(key, []).append(index)
                    for page_sequence, indexes in enumerate(groups.values()):
                        text = sanitize_source_text(" ".join(data["text"][index].strip() for index in indexes)).strip()
                        if not text:
                            continue
                        left = min(data["left"][index] for index in indexes); top = min(data["top"][index] for index in indexes)
                        right = max(data["left"][index] + data["width"][index] for index in indexes); bottom = max(data["top"][index] + data["height"][index] for index in indexes)
                        confidences = [float(data["conf"][index]) for index in indexes if float(data["conf"][index]) >= 0]
                        confidence = (sum(confidences) / len(confidences) / 100) if confidences else .5
                        blocks.append(BlockEvidence(len(blocks), page_sequence, ExtractedBlockType.PARAGRAPH, EvidenceOrigin.OCR_INFERRED, text, text, _page_reference(page_index, reader.pages[page_index]), {"x0": left, "y0": top, "x1": right, "y1": bottom, "coordinate_space": "ocr_pixels"}, source_method="tesseract", confidence=max(0, min(confidence, 1))))
                    ocr_count += 1
                ocr_engine = "tesseract"
        blocks = [BlockEvidence(index, block.page_sequence_number, block.block_type, block.evidence_origin, block.raw_text, block.normalized_text, block.page_reference, block.geometry, block.typography, block.structural_hints, block.source_method, block.table_reference, block.image_reference, block.confidence, block.metadata) for index, block in enumerate(sorted(blocks, key=lambda block: (block.page_reference.get("page_index", 0), block.page_sequence_number, block.sequence_number)))]
        method = ExtractionMethod.PDF_MIXED if native_pages and ocr_count else ExtractionMethod.PDF_OCR if ocr_count else ExtractionMethod.PDF_NATIVE
        return ExtractionEvidence(method, tuple(blocks), len(reader.pages), native_pages=native_pages, ocr_pages=ocr_count, ocr_engine=ocr_engine, warnings=tuple(warnings))


class DocxDocumentExtractor:
    name, version = "python-docx-style-extractor", EXTRACTOR_VERSION

    def supports(self, detected_format) -> bool:
        return detected_format == SourceFormat.DOCX

    def extract(self, source: SourceDocument, profile) -> ExtractionEvidence:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        document = Document(io.BytesIO(source.content))
        blocks: list[BlockEvidence] = []
        for child in document.element.body.iterchildren():
            if child.tag.endswith("}p"):
                paragraph = Paragraph(child, document)
                text = sanitize_source_text(paragraph.text).strip()
                if not text:
                    continue
                style = paragraph.style.name if paragraph.style else ""
                style_lower = style.lower()
                block_type = ExtractedBlockType.TITLE if style_lower == "title" else ExtractedBlockType.HEADING_1 if style_lower.startswith("heading 1") else ExtractedBlockType.HEADING_2 if style_lower.startswith("heading 2") else ExtractedBlockType.HEADING_3 if style_lower.startswith("heading 3") else ExtractedBlockType.LIST_ITEM if "list" in style_lower else ExtractedBlockType.CAPTION if "caption" in style_lower else ExtractedBlockType.PARAGRAPH
                blocks.append(BlockEvidence(len(blocks), len(blocks), block_type, EvidenceOrigin.STYLE_INFERRED if style else EvidenceOrigin.PARSER_DEFAULT, text, re.sub(r"[ \t]+", " ", text), structural_hints={"style_name": style}, source_method="python_docx", confidence=.95 if style else .8))
            elif child.tag.endswith("}tbl"):
                table = Table(child, document)
                table_ref = f"table-{sum(1 for b in blocks if b.block_type == ExtractedBlockType.TABLE) + 1}"
                blocks.append(BlockEvidence(len(blocks), len(blocks), ExtractedBlockType.TABLE, EvidenceOrigin.SOURCE_EXPLICIT, source_method="python_docx", table_reference=table_ref, confidence=.98))
                for row_index, row in enumerate(table.rows):
                    row_text = " | ".join(sanitize_source_text(cell.text).strip() for cell in row.cells)
                    blocks.append(BlockEvidence(len(blocks), len(blocks), ExtractedBlockType.TABLE_ROW, EvidenceOrigin.SOURCE_EXPLICIT, row_text, row_text, structural_hints={"row_index": row_index}, source_method="python_docx", table_reference=table_ref, confidence=.98))
                    for cell_index, cell in enumerate(row.cells):
                        cell_text = sanitize_source_text(cell.text).strip()
                        blocks.append(BlockEvidence(len(blocks), len(blocks), ExtractedBlockType.TABLE_CELL, EvidenceOrigin.SOURCE_EXPLICIT, cell_text, cell_text, structural_hints={"row_index": row_index, "cell_index": cell_index}, source_method="python_docx", table_reference=table_ref, confidence=.98))
        return ExtractionEvidence(ExtractionMethod.DOCX_NATIVE, tuple(blocks), None, native_text_pages=1 if blocks else 0)
