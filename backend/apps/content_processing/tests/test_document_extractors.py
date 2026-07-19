import io
from types import SimpleNamespace
from unittest.mock import patch

from docx import Document

from apps.content_processing.application.extraction_ports import SourceDocument
from apps.content_processing.domain.extraction import ExtractionMethod
from apps.content_processing.infrastructure.document_extraction import (
    DocxDocumentExtractor,
    PdfDocumentExtractor,
)


class FakeBox:
    width = 612
    height = 792


class FakePdfPage:
    mediabox = FakeBox()
    rotation = 0
    images = ()

    def __init__(self, text=""):
        self.text = text

    def extract_text(self, visitor_text=None):
        if self.text and visitor_text:
            visitor_text(self.text, None, [1, 0, 0, 1, 72, 720], {"/BaseFont": "Helvetica"}, 12)
        return self.text


def source(filename, content):
    return SourceDocument("file-1", filename, "", len(content), "checksum", content)


def profile(ocr_pages=()):
    return SimpleNamespace(ocr_pages_recommended=list(ocr_pages), ocr_requirement="not_required")


def test_native_pdf_reports_native_text_pages_separately_from_ocr_pages():
    reader = SimpleNamespace(pages=[FakePdfPage("Native academic text")])

    with patch("pypdf.PdfReader", return_value=reader):
        evidence = PdfDocumentExtractor().extract(source("notes.pdf", b"%PDF"), profile())

    assert evidence.method == ExtractionMethod.PDF_NATIVE
    assert evidence.native_text_pages == 1
    assert evidence.ocr_pages == 0
    assert evidence.page_count == 1


def test_mixed_pdf_reports_native_and_ocr_page_counts_independently():
    reader = SimpleNamespace(pages=[FakePdfPage("Native academic text"), FakePdfPage()])
    ocr_data = {
        "text": ["Scanned", "evidence"],
        "block_num": [1, 1],
        "par_num": [1, 1],
        "line_num": [1, 1],
        "left": [10, 70],
        "top": [20, 20],
        "width": [50, 60],
        "height": [12, 12],
        "conf": ["90", "80"],
    }

    with (
        patch("pypdf.PdfReader", return_value=reader),
        patch("apps.content_processing.infrastructure.document_extraction.shutil.which", return_value="/usr/bin/tool"),
        patch("pdf2image.convert_from_bytes", return_value=[object()]),
        patch("pytesseract.image_to_data", return_value=ocr_data),
    ):
        evidence = PdfDocumentExtractor().extract(source("mixed.pdf", b"%PDF"), profile((1,)))

    assert evidence.method == ExtractionMethod.PDF_MIXED
    assert evidence.native_text_pages == 1
    assert evidence.ocr_pages == 1
    assert evidence.page_count == 2


def test_docx_native_text_page_contract_is_unchanged():
    document = Document()
    document.add_paragraph("A native DOCX paragraph.")
    payload = io.BytesIO()
    document.save(payload)

    evidence = DocxDocumentExtractor().extract(source("notes.docx", payload.getvalue()), profile())

    assert evidence.method == ExtractionMethod.DOCX_NATIVE
    assert evidence.native_text_pages == 1
    assert evidence.ocr_pages == 0
