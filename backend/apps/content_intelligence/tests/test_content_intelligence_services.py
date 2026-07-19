from io import BytesIO
from types import SimpleNamespace
from unittest import TestCase
from unittest.mock import Mock, patch
import zipfile

from django.test import SimpleTestCase
from PIL import Image, ImageDraw
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import ImageReader
from reportlab.pdfgen import canvas

from apps.content_intelligence.application import (
    ConceptExtractionService,
    ConceptCandidateValidator,
    ContentImportDeletionService,
    DocumentTextNormalizationService,
    ExtractionService,
    HeadingNormalizationService,
    ImportService,
    PipelineService,
    SectionDetectionService,
    ValidationService,
)
from apps.content_intelligence.domain.exceptions import ExtractionError
from apps.content_intelligence.domain.models import ContentExtractionResult, ContentImportJob, ParsedDocument, ParsedSection
from apps.content_intelligence.domain.services import ExtractionPayload
from apps.content_intelligence.infrastructure.extraction.adapters import DocxExtractionAdapter, PdfExtractionAdapter


class ContentIntelligenceServiceTests(SimpleTestCase):
    def test_import_service_creates_job_and_publishes_event(self):
        repository = Mock()
        publisher = Mock()
        repository.add.side_effect = lambda job: self._with_id(job)
        service = ImportService(job_repository=repository, event_publisher=publisher)
        processing_job = SimpleNamespace(id="processing-1")

        with patch.object(service, "_create_processing_job_service") as create_processing_job_service:
            with patch.object(service, "_queue_processing_job_service") as queue_processing_job_service:
                create_processing_job_service.return_value.create_or_resolve.return_value = processing_job
                queue_processing_job_service.return_value.queue.return_value = processing_job
                repository.get.side_effect = lambda job_id: SimpleNamespace(
                    id=job_id,
                    format_type=ContentImportJob.FormatType.DOCX,
                    processing_job=processing_job,
                )

                job = service.create_import_job(self._resource("notes.docx"))

        self.assertEqual(job.format_type, ContentImportJob.FormatType.DOCX)
        self.assertEqual(publisher.publish.call_args.args[0].event_name, "content_intelligence.import_started")
        create_processing_job_service.return_value.create_or_resolve.assert_called_once()
        queue_processing_job_service.return_value.queue.assert_called_once_with(processing_job)

    def test_import_service_rejects_unsupported_format(self):
        service = ImportService(job_repository=Mock(), event_publisher=Mock())

        with self.assertRaisesMessage(Exception, "Only PDF and DOCX uploads are supported."):
            service.create_import_job(self._resource("notes.txt"))

    def test_extraction_service_requests_ocr_when_text_is_insufficient(self):
        result_repository = Mock()
        storage_provider = Mock()
        storage_provider.download.side_effect = [BytesIO(b"pdf"), BytesIO(b"ocr")]
        pdf_adapter = Mock()
        pdf_adapter.extract.return_value = ExtractionPayload("", "", "pdf_text", False, metadata={})
        ocr_service = Mock()
        ocr_service.extract_text.return_value = ExtractionPayload("ocr text", "ocr text", "pdf_ocr_fallback", True, metadata={})
        result_repository.get_for_job.return_value = None
        result_repository.save_result.side_effect = lambda result: self._with_id(result)
        publisher = Mock()
        service = ExtractionService(
            result_repository=result_repository,
            storage_provider=storage_provider,
            pdf_adapter=pdf_adapter,
            ocr_service=ocr_service,
            event_publisher=publisher,
        )

        job = ContentImportJob(
            id="job-1",
            learning_resource=self._resource("scan.pdf"),
            stored_file=SimpleNamespace(id="file-1", stored_filename="scan.pdf"),
            format_type=ContentImportJob.FormatType.PDF,
        )
        result = service.extract(job)

        self.assertTrue(result.ocr_requested)
        self.assertTrue(result.ocr_used)
        self.assertEqual([call.args[0].event_name for call in publisher.publish.call_args_list], [
            "content_intelligence.ocr_requested",
            "content_intelligence.ocr_completed",
            "content_intelligence.extraction_completed",
        ])

    def test_extraction_service_uses_fresh_stream_for_primary_extraction(self):
        result_repository = Mock()
        storage_provider = Mock()
        source_stream = BytesIO(b"sample-bytes")
        source_stream.read()
        storage_provider.download.return_value = source_stream
        pdf_adapter = Mock()

        def extract(stream):
            self.assertEqual(stream.tell(), 0)
            self.assertEqual(stream.read(), b"sample-bytes")
            return ExtractionPayload("This is valid PDF text.", "This is valid PDF text.", "pdf_text", True, metadata={})

        pdf_adapter.extract.side_effect = extract
        result_repository.get_for_job.return_value = None
        result_repository.save_result.side_effect = lambda result: self._with_id(result)

        service = ExtractionService(
            result_repository=result_repository,
            storage_provider=storage_provider,
            pdf_adapter=pdf_adapter,
            event_publisher=Mock(),
        )
        job = ContentImportJob(
            id="job-1",
            learning_resource=self._resource("fresh.pdf"),
            stored_file=SimpleNamespace(id="file-1", stored_filename="fresh.pdf"),
            format_type=ContentImportJob.FormatType.PDF,
        )

        result = service.extract(job)

        self.assertTrue(result.sufficient_text)
        pdf_adapter.extract.assert_called_once()

    def test_native_pdf_extraction_avoids_ocr_when_text_is_sufficient(self):
        result_repository = Mock()
        storage_provider = Mock()
        storage_provider.download.return_value = BytesIO(b"pdf")
        pdf_adapter = Mock()
        pdf_adapter.extract.return_value = ExtractionPayload(
            "This clean PDF has enough text to pass.",
            "This clean PDF has enough text to pass.",
            "pdf_text",
            True,
            metadata={},
        )
        ocr_service = Mock()
        result_repository.get_for_job.return_value = None
        result_repository.save_result.side_effect = lambda result: self._with_id(result)

        service = ExtractionService(
            result_repository=result_repository,
            storage_provider=storage_provider,
            pdf_adapter=pdf_adapter,
            ocr_service=ocr_service,
            event_publisher=Mock(),
        )
        job = ContentImportJob(
            id="job-1",
            learning_resource=self._resource("clean.pdf"),
            stored_file=SimpleNamespace(id="file-1", stored_filename="clean.pdf"),
            format_type=ContentImportJob.FormatType.PDF,
        )

        result = service.extract(job)

        self.assertTrue(result.sufficient_text)
        ocr_service.extract_text.assert_not_called()

    def test_short_but_valid_documents_are_not_rejected(self):
        result_repository = Mock()
        storage_provider = Mock()
        storage_provider.download.return_value = BytesIO(b"pdf")
        pdf_adapter = Mock()
        text = "Short valid study text."
        pdf_adapter.extract.return_value = ExtractionPayload(text, text, "pdf_text", True, metadata={})
        result_repository.get_for_job.return_value = None
        result_repository.save_result.side_effect = lambda result: self._with_id(result)

        service = ExtractionService(
            result_repository=result_repository,
            storage_provider=storage_provider,
            pdf_adapter=pdf_adapter,
            event_publisher=Mock(),
        )
        job = ContentImportJob(
            id="job-1",
            learning_resource=self._resource("short.pdf"),
            stored_file=SimpleNamespace(id="file-1", stored_filename="short.pdf"),
            format_type=ContentImportJob.FormatType.PDF,
        )

        result = service.extract(job)

        self.assertEqual(result.normalized_text, text)
        self.assertTrue(result.sufficient_text)

    def test_extraction_service_sanitizes_nul_bytes_before_persistence(self):
        result_repository = Mock()
        storage_provider = Mock()
        storage_provider.download.return_value = BytesIO(b"pdf")
        pdf_adapter = Mock()
        pdf_adapter.extract.return_value = ExtractionPayload(
            "Hello\x00 world",
            "Hello\x00 world",
            "pdf_text",
            True,
            metadata={"note": "bad\x00text"},
        )
        result_repository.get_for_job.return_value = None
        result_repository.save_result.side_effect = lambda result: self._with_id(result)

        service = ExtractionService(
            result_repository=result_repository,
            storage_provider=storage_provider,
            pdf_adapter=pdf_adapter,
            event_publisher=Mock(),
        )
        job = ContentImportJob(
            id="job-1",
            learning_resource=self._resource("nul.pdf"),
            stored_file=SimpleNamespace(id="file-1", stored_filename="nul.pdf"),
            format_type=ContentImportJob.FormatType.PDF,
        )

        result = service.extract(job)

        self.assertEqual(result.extracted_text, "Hello world")
        self.assertEqual(result.normalized_text, "Hello world")
        self.assertEqual(result.metadata["note"], "badtext")
        self.assertEqual(result.metadata["nul_bytes_removed"], 3)

    def test_image_only_pdf_requests_ocr_when_native_extraction_is_insufficient(self):
        result_repository = Mock()
        storage_provider = Mock()
        storage_provider.download.return_value = BytesIO(self._image_only_pdf_bytes())
        ocr_service = Mock()

        def extract_text(stream, format_type):
            self.assertEqual(stream.tell(), 0)
            self.assertEqual(format_type, ContentImportJob.FormatType.PDF)
            return ExtractionPayload(
                "Recovered OCR study text for the scanned handout.",
                "Recovered OCR study text for the scanned handout.",
                "pdf_ocr_fallback",
                True,
                metadata={"warning": None},
            )

        ocr_service.extract_text.side_effect = extract_text
        result_repository.get_for_job.return_value = None
        result_repository.save_result.side_effect = lambda result: self._with_id(result)

        service = ExtractionService(
            result_repository=result_repository,
            storage_provider=storage_provider,
            pdf_adapter=PdfExtractionAdapter(),
            ocr_service=ocr_service,
            event_publisher=Mock(),
        )
        job = ContentImportJob(
            id="job-1",
            learning_resource=self._resource("scan.pdf"),
            stored_file=SimpleNamespace(id="file-1", original_filename="scan.pdf", stored_filename="scan.pdf"),
            format_type=ContentImportJob.FormatType.PDF,
        )

        result = service.extract(job)

        self.assertTrue(result.ocr_requested)
        self.assertTrue(result.ocr_used)
        ocr_service.extract_text.assert_called_once()

    def test_ocr_unavailability_surfaces_precise_failure_reason(self):
        result_repository = Mock()
        storage_provider = Mock()
        storage_provider.download.return_value = BytesIO(self._image_only_pdf_bytes())
        ocr_service = Mock()
        ocr_service.extract_text.return_value = ExtractionPayload(
            "",
            "",
            "pdf_ocr_fallback",
            False,
            metadata={"warning": "OCR engine unavailable; PDF rasterizer unavailable"},
        )

        service = ExtractionService(
            result_repository=result_repository,
            storage_provider=storage_provider,
            pdf_adapter=PdfExtractionAdapter(),
            ocr_service=ocr_service,
            event_publisher=Mock(),
        )
        job = ContentImportJob(
            id="job-1",
            learning_resource=self._resource("scan.pdf"),
            stored_file=SimpleNamespace(id="file-1", original_filename="scan.pdf", stored_filename="scan.pdf"),
            format_type=ContentImportJob.FormatType.PDF,
        )

        with self.assertRaises(ExtractionError) as exc:
            service.extract(job)

        self.assertEqual(
            exc.exception.details["failure_reason"],
            "OCR engine unavailable; PDF rasterizer unavailable",
        )

    def test_section_detection_service_handles_headings(self):
        service = SectionDetectionService(event_publisher=Mock())
        parsed_document = ParsedDocument(
            id="document-1",
            import_job=ContentImportJob(id="job-1", learning_resource=self._resource(), format_type=ContentImportJob.FormatType.PDF),
            normalized_text="Chapter 1 Foundations\n\nBody one\n\nChapter 2 Applications\n\nBody two",
            format_type=ContentImportJob.FormatType.PDF,
            extraction_method="pdf_text",
        )

        sections = service.detect_sections(parsed_document)

        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].heading, "Chapter 1 Foundations")

    def test_concept_extraction_service_builds_candidates(self):
        service = ConceptExtractionService(event_publisher=Mock())
        section = Mock()
        section.id = "section-1"
        section.heading = "Chapter 1 Foundations"
        section.body_text = "A concept paragraph.\n\nAnother concept paragraph."
        section.parsed_document = Mock(import_job_id="job-1")

        concepts = service.extract_concepts(section)

        self.assertEqual(len(concepts), 2)
        self.assertEqual(concepts[0].sequence_number, 1)

    def test_reprocessing_replaces_obsolete_academic_concepts(self):
        job_repository = Mock()
        parsed_document_repository = Mock()
        validation_repository = Mock()
        run_repository = Mock()
        extraction_service = Mock()
        section_detection_service = Mock()
        concept_extraction_service = Mock()
        validation_service = Mock()
        learning_content_service = Mock()
        publisher = Mock()

        resource = self._resource("resource.pdf")
        job = ContentImportJob(id="job-1", learning_resource=resource, format_type=ContentImportJob.FormatType.PDF)
        extraction_result = ContentExtractionResult(
            import_job=job,
            normalized_text="Chapter 1: Cell Structure\n\nCells contain organelles.",
            extraction_method="pdf_text",
            sufficient_text=True,
            char_count=500,
            page_count=2,
        )
        parsed_document = ParsedDocument(id="doc-1", import_job=job, normalized_text="Cell Structure\n\nCells contain organelles.", format_type=job.format_type, extraction_method="pdf_text")
        accepted_candidate = SimpleNamespace(
            title="Cell Structure",
            description="Cells contain organelles.",
            learning_objective="Understand cell structure",
            sequence_number=1,
            confidence=0.91,
            metadata={"decision": "accepted", "rejection_reasons": [], "title_semantic_key": "cell structure"},
        )
        section = Mock()
        section.heading = "Chapter 1: Cell Structure"
        section.sequence_number = 1
        section.body_text = "Cells contain organelles."
        section.metadata = {"semantic_title": "Cell Structure", "section_classification": "academic_content"}
        section.concept_candidates.all.return_value.order_by.return_value = [accepted_candidate]
        extraction_service.extract.return_value = extraction_result
        parsed_document_repository.get_for_job.return_value = None
        parsed_document_repository.save_document.return_value = parsed_document
        section_detection_service.detect_sections.return_value = [section]
        parsed_document_repository.replace_sections.return_value = [section]
        concept_extraction_service.extract_concepts.return_value = [accepted_candidate]
        parsed_document_repository.replace_concepts.return_value = [accepted_candidate]
        validation_service.validate.return_value = []
        validation_repository.replace_for_job.return_value = []
        run_repository.add.side_effect = lambda run: run
        run_repository.save.side_effect = lambda run: run
        job_repository.save.side_effect = lambda saved_job: saved_job
        existing_section = Mock(sequence_number=1)
        existing_concept_kept = Mock(sequence_number=1)
        existing_concept_removed = Mock(sequence_number=2)
        learning_content_service.list_sections.return_value = [existing_section]
        learning_content_service.update_section.return_value = existing_section
        learning_content_service.list_concepts.return_value = [existing_concept_kept, existing_concept_removed]

        service = PipelineService(
            job_repository=job_repository,
            parsed_document_repository=parsed_document_repository,
            validation_repository=validation_repository,
            pipeline_run_repository=run_repository,
            extraction_service=extraction_service,
            section_detection_service=section_detection_service,
            concept_extraction_service=concept_extraction_service,
            validation_service=validation_service,
            learning_content_service=learning_content_service,
            event_publisher=publisher,
        )

        completed = service.run_pipeline(job)

        self.assertEqual(completed.status, ContentImportJob.Status.COMPLETED)
        learning_content_service.update_concept.assert_called_once()
        learning_content_service.create_concept.assert_not_called()
        learning_content_service.delete_concept.assert_called_once_with(existing_concept_removed)

    def test_pipeline_service_runs_and_populates_academic_content(self):
        job_repository = Mock()
        parsed_document_repository = Mock()
        validation_repository = Mock()
        run_repository = Mock()
        extraction_service = Mock()
        section_detection_service = Mock()
        concept_extraction_service = Mock()
        validation_service = Mock()
        learning_content_service = Mock()
        publisher = Mock()

        resource = self._resource("resource.pdf")
        job = ContentImportJob(id="job-1", learning_resource=resource, format_type=ContentImportJob.FormatType.PDF)
        extraction_result = ContentExtractionResult(
            import_job=job,
            normalized_text="Chapter 1\n\nBody",
            extraction_method="pdf_text",
            sufficient_text=True,
            char_count=500,
            page_count=2,
        )
        parsed_document = ParsedDocument(id="doc-1", import_job=job, normalized_text="Chapter 1\n\nBody", format_type=job.format_type, extraction_method="pdf_text")
        section = Mock()
        section.heading = "Chapter 1"
        section.sequence_number = 1
        section.body_text = "Body"
        section.concept_candidates.all.return_value.order_by.return_value = [SimpleNamespace(sequence_number=1, title="Concept", description="Body", learning_objective="Understand concept")]
        extraction_service.extract.return_value = extraction_result
        parsed_document_repository.get_for_job.return_value = None
        parsed_document_repository.save_document.return_value = parsed_document
        section_detection_service.detect_sections.return_value = [section]
        parsed_document_repository.replace_sections.return_value = [section]
        concept_extraction_service.extract_concepts.return_value = [SimpleNamespace()]
        parsed_document_repository.replace_concepts.return_value = [SimpleNamespace(confidence=0.8)]
        validation_service.validate.return_value = []
        validation_repository.replace_for_job.return_value = []
        run_repository.add.side_effect = lambda run: run
        run_repository.save.side_effect = lambda run: run
        job_repository.save.side_effect = lambda saved_job: saved_job
        learning_content_service.list_sections.return_value = []
        learning_content_service.create_section.return_value = Mock()
        learning_content_service.list_concepts.return_value = []

        service = PipelineService(
            job_repository=job_repository,
            parsed_document_repository=parsed_document_repository,
            validation_repository=validation_repository,
            pipeline_run_repository=run_repository,
            extraction_service=extraction_service,
            section_detection_service=section_detection_service,
            concept_extraction_service=concept_extraction_service,
            validation_service=validation_service,
            learning_content_service=learning_content_service,
            event_publisher=publisher,
        )

        completed = service.run_pipeline(job)

        self.assertEqual(completed.status, ContentImportJob.Status.COMPLETED)
        self.assertTrue(any(call.args[0].event_name == "content_intelligence.import_completed" for call in publisher.publish.call_args_list))

    def test_pipeline_marks_job_failed_when_pdf_and_ocr_extraction_are_insufficient(self):
        job_repository = Mock()
        parsed_document_repository = Mock()
        validation_repository = Mock()
        run_repository = Mock()
        section_detection_service = Mock()
        concept_extraction_service = Mock()
        validation_service = Mock()
        learning_content_service = Mock()
        publisher = Mock()

        storage_provider = Mock()
        storage_provider.download.side_effect = [BytesIO(b"%PDF-\x00binary"), BytesIO(b"%PDF-\x00binary")]
        pdf_adapter = Mock()
        pdf_adapter.extract.return_value = ExtractionPayload("", "", "pdf_text", False, metadata={})
        ocr_service = Mock()
        ocr_service.extract_text.return_value = ExtractionPayload("", "", "pdf_ocr_fallback", False, metadata={})
        extraction_service = ExtractionService(
            result_repository=Mock(get_for_job=Mock(return_value=None), save_result=Mock(side_effect=lambda result: result)),
            storage_provider=storage_provider,
            pdf_adapter=pdf_adapter,
            ocr_service=ocr_service,
            event_publisher=publisher,
        )

        resource = self._resource("resource.pdf")
        job = ContentImportJob(
            id="job-1",
            learning_resource=resource,
            stored_file=SimpleNamespace(id="file-1", stored_filename="resource.pdf"),
            format_type=ContentImportJob.FormatType.PDF,
        )
        run_repository.add.side_effect = lambda run: run
        run_repository.save.side_effect = lambda run: run
        job_repository.save.side_effect = lambda saved_job: saved_job

        service = PipelineService(
            job_repository=job_repository,
            parsed_document_repository=parsed_document_repository,
            validation_repository=validation_repository,
            pipeline_run_repository=run_repository,
            extraction_service=extraction_service,
            section_detection_service=section_detection_service,
            concept_extraction_service=concept_extraction_service,
            validation_service=validation_service,
            learning_content_service=learning_content_service,
            event_publisher=publisher,
        )

        with self.assertRaises(ExtractionError):
            service.run_pipeline(job)

        self.assertEqual(job.status, ContentImportJob.Status.FAILED)
        self.assertTrue(any(call.args[0].event_name == "content_intelligence.import_failed" for call in publisher.publish.call_args_list))

    def test_deletion_service_retires_academic_and_retrieval_content_and_publishes_events(self):
        learning_content_service = Mock()
        learning_resource_service = Mock()
        storage_service = Mock()
        retrieval_retirement_service = Mock()
        publisher = Mock()

        section = SimpleNamespace(id="section-1")
        concept = SimpleNamespace(id="concept-1")
        learning_content_service.list_sections.return_value = [section]
        learning_content_service.list_concepts.return_value = [concept]

        resource = self._resource("delete.pdf")
        resource.status = "active"
        resource.id = "resource-1"
        job = ContentImportJob(
            id="job-1",
            learning_resource=resource,
            stored_file=resource.stored_file,
            format_type=ContentImportJob.FormatType.PDF,
            status=ContentImportJob.Status.FAILED,
        )
        job.delete = Mock()

        service = ContentImportDeletionService(
            learning_content_service=learning_content_service,
            learning_resource_service=learning_resource_service,
            storage_service=storage_service,
            retrieval_retirement_service=retrieval_retirement_service,
            event_publisher=publisher,
        )

        atomic = Mock()
        atomic.return_value.__enter__ = Mock(return_value=None)
        atomic.return_value.__exit__ = Mock(return_value=False)

        with patch("apps.content_intelligence.application.deletion_service.transaction.atomic", atomic):
            result = service.delete_import(job)

        self.assertEqual(result.deleted_sections, 1)
        self.assertEqual(result.deleted_concepts, 1)
        retrieval_retirement_service.retire.assert_called_once_with(resource)
        learning_content_service.archive_concept.assert_called_once_with(concept)
        learning_content_service.archive_section.assert_called_once_with(section)
        learning_resource_service.update_resource.assert_called_once_with(resource, stored_file=None)
        learning_resource_service.archive_resource.assert_called_once_with(resource)
        storage_service.delete_file_contents.assert_called_once_with(resource.stored_file)
        self.assertEqual(
            [call.args[0].event_name for call in publisher.publish.call_args_list],
            ["content_intelligence.deletion_requested", "content_intelligence.deleted"],
        )

    def test_deletion_service_rejects_processing_jobs(self):
        service = ContentImportDeletionService(
            learning_content_service=Mock(),
            learning_resource_service=Mock(),
            storage_service=Mock(),
            event_publisher=Mock(),
        )
        job = ContentImportJob(
            id="job-1",
            learning_resource=self._resource("processing.pdf"),
            stored_file=self._resource("processing.pdf").stored_file,
            format_type=ContentImportJob.FormatType.PDF,
            status=ContentImportJob.Status.PROCESSING,
        )

        atomic = Mock()
        atomic.return_value.__enter__ = Mock(return_value=None)
        atomic.return_value.__exit__ = Mock(return_value=False)

        with patch("apps.content_intelligence.application.deletion_service.transaction.atomic", atomic):
            with self.assertRaisesMessage(Exception, "cannot be deleted safely yet"):
                service.delete_import(job)

    def test_deletion_service_allows_processing_jobs_when_processing_aggregate_exists(self):
        learning_content_service = Mock()
        learning_resource_service = Mock()
        storage_service = Mock()
        publisher = Mock()

        resource = self._resource("processing.pdf")
        resource.id = "resource-1"
        processing_job = SimpleNamespace(id="processing-1", status="active")
        job = ContentImportJob(
            id="job-1",
            learning_resource=resource,
            stored_file=resource.stored_file,
            format_type=ContentImportJob.FormatType.PDF,
            status=ContentImportJob.Status.PROCESSING,
        )
        job.processing_job = processing_job
        job.delete = Mock()
        learning_content_service.list_sections.return_value = []

        service = ContentImportDeletionService(
            learning_content_service=learning_content_service,
            learning_resource_service=learning_resource_service,
            storage_service=storage_service,
            event_publisher=publisher,
            retrieval_retirement_service=Mock(),
        )

        atomic = Mock()
        atomic.return_value.__enter__ = Mock(return_value=None)
        atomic.return_value.__exit__ = Mock(return_value=False)

        with patch("apps.content_intelligence.application.deletion_service.transaction.atomic", atomic):
            with patch.object(service, "_cancel_processing_service") as cancel_service:
                with patch.object(service, "_delete_processing_service") as delete_service:
                    cancel_service.return_value.cancel.return_value = processing_job
                    result = service.delete_import(job)

        self.assertEqual(result.content_import_job_id, "job-1")
        cancel_service.return_value.cancel.assert_called_once_with(processing_job)
        delete_service.return_value.mark_deleted.assert_called_once_with(processing_job)

    def _resource(self, filename="resource.pdf"):
        return SimpleNamespace(
            id="resource-1",
            title="Resource",
            stored_file=SimpleNamespace(id="file-1", original_filename=filename, stored_filename=filename),
        )

    def _with_id(self, obj):
        obj.id = getattr(obj, "id", "generated-id")
        return obj

    def _image_only_pdf_bytes(self) -> bytes:
        image = Image.new("RGB", (600, 180), color="white")
        draw = ImageDraw.Draw(image)
        draw.text((24, 72), "Scanned biology notes", fill="black")

        image_buffer = BytesIO()
        image.save(image_buffer, format="PNG")
        image_buffer.seek(0)

        pdf_buffer = BytesIO()
        pdf_canvas = canvas.Canvas(pdf_buffer, pagesize=letter)
        pdf_canvas.drawImage(ImageReader(image_buffer), 72, 540, width=420, height=126)
        pdf_canvas.showPage()
        pdf_canvas.save()
        return pdf_buffer.getvalue()


class ContentIntelligenceAdapterTests(TestCase):
    def test_real_pdf_parser_extracts_known_text(self):
        adapter = PdfExtractionAdapter()

        payload = adapter.extract(BytesIO(self._pdf_bytes("Cell structure notes help explain organelles and membranes.")))

        self.assertTrue(payload.sufficient_text)
        self.assertIn("Cell structure notes", payload.normalized_text)
        self.assertEqual(payload.metadata["parser_library"], "pypdf")
        self.assertTrue(payload.metadata["dependency_available"])

    def test_real_docx_parser_extracts_known_text(self):
        adapter = DocxExtractionAdapter()

        payload = adapter.extract(BytesIO(self._docx_bytes(
            "This DOCX contains enough valid text for extraction.",
            table_rows=[("Topic", "Cell membrane")],
        )))

        self.assertTrue(payload.sufficient_text)
        self.assertIn("valid text", payload.normalized_text)
        self.assertIn("Cell membrane", payload.normalized_text)
        self.assertEqual(payload.metadata["parser_library"], "python-docx")

    def test_clean_pdf_produces_sufficient_extracted_text(self):
        adapter = PdfExtractionAdapter()
        payload = adapter.extract(BytesIO(self._pdf_bytes("This PDF contains enough clean text.")))

        self.assertTrue(payload.sufficient_text)
        self.assertIn("clean text", payload.normalized_text)

    def test_clean_docx_produces_sufficient_extracted_text(self):
        adapter = DocxExtractionAdapter()
        payload = adapter.extract(BytesIO(self._docx_bytes("This DOCX contains enough valid text.")))

        self.assertTrue(payload.sufficient_text)
        self.assertIn("valid text", payload.normalized_text)

    def test_pdf_adapter_does_not_return_raw_binary_fallback_text(self):
        adapter = PdfExtractionAdapter()
        fake_reader = Mock(side_effect=RuntimeError("pdf parse failed"))

        with patch.dict("sys.modules", {"pypdf": SimpleNamespace(PdfReader=fake_reader)}):
            payload = adapter.extract(BytesIO(b"%PDF-\x00\x01binary"))

        self.assertEqual(payload.extracted_text, "")
        self.assertEqual(payload.normalized_text, "")
        self.assertFalse(payload.sufficient_text)
        self.assertFalse(payload.metadata["parser_succeeded"])

    def test_invalid_docx_raises_extraction_error(self):
        adapter = DocxExtractionAdapter()

        with self.assertRaises(ExtractionError):
            adapter.extract(BytesIO(b"PK\x03\x04\x00\x00binary"))

    def _pdf_bytes(self, text: str) -> bytes:
        buffer = BytesIO()
        pdf_canvas = canvas.Canvas(buffer, pagesize=letter)
        pdf_canvas.drawString(72, 720, text)
        pdf_canvas.showPage()
        pdf_canvas.save()
        return buffer.getvalue()

    def _docx_bytes(self, paragraph: str, table_rows=None) -> bytes:
        table_rows = table_rows or []
        document = Document()
        document.add_paragraph(paragraph)
        if table_rows:
            table = document.add_table(rows=0, cols=len(table_rows[0]))
            for row_values in table_rows:
                row = table.add_row().cells
                for index, value in enumerate(row_values):
                    row[index].text = value
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()


class ContentIntelligenceSemanticParsingTests(SimpleTestCase):
    def setUp(self):
        self.normalizer = DocumentTextNormalizationService()
        self.heading_normalizer = HeadingNormalizationService()
        self.validator = ConceptCandidateValidator()
        self.section_detection_service = SectionDetectionService(event_publisher=Mock())
        self.concept_extraction_service = ConceptExtractionService(event_publisher=Mock())

    def test_date_rejection(self):
        assessment = self.validator.validate("September 19", "September 19", self._section())

        self.assertEqual(assessment.decision, "rejected")
        self.assertIn("date_like", assessment.rejection_reasons)

    def test_contents_rejection(self):
        assessment = self.validator.validate("ii Contents Chapter One", "ii Contents Chapter One", self._section())

        self.assertEqual(assessment.decision, "rejected")
        self.assertIn("navigation_text", assessment.rejection_reasons)

    def test_broken_token_rejection(self):
        assessment = self.validator.validate("Concept 2Available", "Concept 2Available", self._section())

        self.assertEqual(assessment.decision, "rejected")
        self.assertIn("malformed_tokenization", assessment.rejection_reasons)

    def test_generic_heading_rejection(self):
        assessment = self.validator.validate("Chapter 1", "Chapter 1", self._section())

        self.assertEqual(assessment.decision, "rejected")
        self.assertIn("generic_structure", assessment.rejection_reasons)

    def test_meaningful_heading_acceptance(self):
        section = self._section(heading="Chapter 1: Cell Structure", body_text="Cells contain specialized organelles that support life.")
        concepts = self.concept_extraction_service.extract_concepts(section)
        accepted = [concept for concept in concepts if concept.metadata.get("decision") == "accepted"]

        self.assertTrue(any(concept.title == "Cell Structure" for concept in accepted))

    def test_toc_duplicate_prefers_body_occurrence(self):
        raw_text = "\n".join(
            [
                "Contents",
                "Cell Structure ..... 7",
                "",
                "Chapter 1: Cell Structure",
                "",
                "Cells contain membranes and organelles that organize the cell.",
            ]
        )
        normalized = self.normalizer.normalize(raw_text)
        parsed_document = ParsedDocument(
            id="doc-1",
            import_job=ContentImportJob(id="job-1", learning_resource=self._resource(), format_type=ContentImportJob.FormatType.PDF),
            normalized_text=normalized.cleaned_text,
            format_type=ContentImportJob.FormatType.PDF,
            extraction_method="pdf_text",
        )

        sections = self.section_detection_service.detect_sections(parsed_document)
        concepts = self.concept_extraction_service.extract_concepts(sections[0])
        accepted_titles = [concept.title for concept in concepts if concept.metadata.get("decision") == "accepted"]

        self.assertEqual(accepted_titles, ["Cell Structure"])

    def test_roman_numeral_page_marker_rejection(self):
        assessment = self.validator.validate("xii", "xii", self._section())

        self.assertEqual(assessment.decision, "rejected")
        self.assertIn("page_marker", assessment.rejection_reasons)

    def test_copyright_metadata_is_front_matter(self):
        classification = self.normalizer.classify_line("Copyright 2024 Abbot Study Press")

        self.assertEqual(classification, "front_matter")

    def test_short_valid_concept_is_accepted(self):
        assessment = self.validator.validate("Mitosis", "Mitosis is the process by which one cell divides into two cells.", self._section())

        self.assertEqual(assessment.decision, "accepted")
        self.assertEqual(assessment.normalized_title, "Mitosis")

    def test_duplicate_page_headers_removed_before_concept_extraction(self):
        raw_text = "\n".join(
            [
                "Biology 101",
                "Chapter 1: Cell Structure",
                "",
                "Cells contain organelles and membranes.",
                "",
                "Biology 101",
                "Chapter 2: Mitosis",
                "",
                "Mitosis is the process of cell division.",
            ]
        )

        normalized = self.normalizer.normalize(raw_text)

        self.assertNotIn("Biology 101", normalized.cleaned_text)
        self.assertGreaterEqual(normalized.metadata["repeated_headers_removed"], 1)

    def test_malformed_ocr_line_is_rejected_with_diagnostic_reason(self):
        assessment = self.validator.validate("ii Contents Chapter One", "ii Contents Chapter One", self._section())

        self.assertIn(assessment.decision, {"manual_review", "rejected"})
        self.assertTrue(assessment.rejection_reasons)

    def test_synthetic_fallback_title_does_not_generate_concept_candidate(self):
        section = self._section(
            heading="Imported Content",
            body_text="Cells contain organelles that support structure and energy use.",
            metadata={"section_classification": "academic_content", "section_origin": "synthetic_fallback"},
        )

        concepts = self.concept_extraction_service.extract_concepts(section)

        self.assertFalse(any(concept.title == "Imported Content" for concept in concepts))
        self.assertEqual(section.metadata.get("synthetic_section_titles_skipped"), 1)

    def test_large_fallback_body_with_semantic_headings_extracts_body_concepts(self):
        section = self._section(
            heading="Imported Content",
            body_text=(
                "Cell Structure\nCells contain specialized organelles that support life.\n\n"
                "Photosynthesis\nPlants convert light energy into chemical energy."
            ),
            metadata={"section_classification": "academic_content", "section_origin": "synthetic_fallback"},
        )

        concepts = self.concept_extraction_service.extract_concepts(section)
        accepted_titles = [concept.title for concept in concepts if concept.metadata.get("decision") == "accepted"]

        self.assertEqual(accepted_titles, ["Cell Structure", "Photosynthesis"])
        self.assertFalse(any(concept.title == "Imported Content" for concept in concepts))

    def test_large_fallback_body_without_structure_keeps_fallback_section_without_fake_concept(self):
        service = ValidationService(event_publisher=Mock())
        section = self._section(
            heading="Imported Content",
            body_text=" ".join(["cells"] * 3000),
            metadata={
                "section_classification": "academic_content",
                "section_origin": "synthetic_fallback",
                "subdivision_attempted": True,
                "inferred_subsections_found": 0,
            },
        )

        concepts = self.concept_extraction_service.extract_concepts(section)
        findings = service.validate(
            job=ContentImportJob(id="job-1", learning_resource=self._resource(), format_type=ContentImportJob.FormatType.PDF),
            sections=[section],
            concepts=concepts,
            extracted_char_count=len(section.body_text),
        )

        self.assertFalse(any(concept.title == "Imported Content" for concept in concepts))
        self.assertTrue(any(f.finding_type == "fallback_structure_uncertain" for f in findings))

    def test_pdf_heading_like_lines_enable_secondary_subdivision(self):
        paragraph = "Cells contain specialized organelles that support life."
        raw_text = (
            "Imported Content\n"
            + "\n".join(
                [
                    "Cell Structure",
                    paragraph,
                    "",
                    "Photosynthesis",
                    "Plants convert light energy into chemical energy.",
                ]
            )
            + "\n\n"
        )
        parsed_document = ParsedDocument(
            id="doc-1",
            import_job=ContentImportJob(id="job-1", learning_resource=self._resource(), format_type=ContentImportJob.FormatType.PDF),
            normalized_text=raw_text * 200,
            format_type=ContentImportJob.FormatType.PDF,
            extraction_method="pdf_text",
            metadata={},
        )

        sections = self.section_detection_service.detect_sections(parsed_document)

        self.assertTrue(any((section.metadata or {}).get("section_origin") == "inferred_heading" for section in sections))
        self.assertTrue(parsed_document.metadata.get("secondary_subdivision_attempted"))

    def test_large_detected_section_is_subdivided_by_internal_chapter_headings(self):
        paragraph = "Economics studies how people allocate scarce resources across competing uses."
        parsed_document = ParsedDocument(
            id="doc-1",
            import_job=ContentImportJob(id="job-1", learning_resource=self._resource(), format_type=ContentImportJob.FormatType.PDF),
            normalized_text=(
                "Introduction to Economics\n\n"
                + "\n".join(
                    [
                        "Chapter 1: Scarcity",
                        paragraph,
                        "",
                        "Chapter 2: Demand",
                        paragraph,
                        "",
                        "Chapter 3: Supply",
                        paragraph,
                    ]
                )
                * 200
            ),
            format_type=ContentImportJob.FormatType.PDF,
            extraction_method="pdf_text",
            metadata={},
        )

        sections = self.section_detection_service.detect_sections(parsed_document)

        self.assertGreaterEqual(len(sections), 3)
        self.assertTrue(any(section.heading == "Chapter 1: Scarcity" for section in sections))
        self.assertFalse(any(section.heading == "Introduction to Economics" and len(section.body_text) > 12000 for section in sections))

    def _section(self, heading="Imported Content", body_text="Cells contain organelles that support structure and energy use.", metadata=None):
        return ParsedSection(
            heading=heading,
            sequence_number=1,
            parsed_document=SimpleNamespace(id="doc-1", import_job_id="job-1"),
            body_text=body_text,
            metadata=metadata or {"section_classification": "academic_content"},
        )

    def _resource(self):
        return SimpleNamespace(id="resource-1", title="Biology Notes")
