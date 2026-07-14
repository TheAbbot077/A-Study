from __future__ import annotations

import io
import re
import zipfile

from apps.content_intelligence.domain.exceptions import ExtractionError
from apps.content_intelligence.domain.services import ExtractionPayload

MIN_MEANINGFUL_CHARACTERS = 20


def meaningful_character_count(value: str) -> int:
    return len(re.findall(r"[A-Za-z0-9]", value))


class PdfExtractionAdapter:
    PARSER_LIBRARY = "pypdf"

    def extract(self, file_obj) -> ExtractionPayload:
        file_obj.seek(0)
        data = file_obj.read()
        signature_hex = data[:4].hex()
        text = ""
        page_count = None
        used_parser = False
        parser_error: str | None = None
        dependency_available = False
        structural_failure: str | None = None
        encrypted = False
        parser_exception_class: str | None = None
        if not data.startswith(b"%PDF"):
            structural_failure = "invalid_pdf_signature"
            parser_error = "Invalid PDF signature."

        try:
            from pypdf import PdfReader  # type: ignore

            dependency_available = True
            if structural_failure is None:
                reader = PdfReader(io.BytesIO(data))
                encrypted = bool(getattr(reader, "is_encrypted", False))
                if encrypted:
                    structural_failure = "encrypted_pdf"
                    parser_error = "Encrypted PDF could not be opened for extraction."
                else:
                    page_count = len(reader.pages)
                    text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
                    used_parser = True
        except Exception as exc:
            text = ""
            parser_exception_class = exc.__class__.__name__
            parser_error = f"{exc.__class__.__name__}: {exc}"
            structural_failure = structural_failure or "malformed_pdf"
        if not dependency_available:
            structural_failure = structural_failure or "parser_dependency_unavailable"

        normalized = self._normalize(text)
        meaningful_count = meaningful_character_count(normalized)
        if (
            dependency_available
            and structural_failure is None
            and used_parser
            and meaningful_count == 0
        ):
            structural_failure = "no_embedded_text"
        return ExtractionPayload(
            extracted_text=text,
            normalized_text=normalized,
            extraction_method="pdf_text",
            sufficient_text=meaningful_count >= MIN_MEANINGFUL_CHARACTERS,
            page_count=page_count,
            metadata={
                "adapter": "pdf",
                "parser_library": self.PARSER_LIBRARY,
                "dependency_available": dependency_available,
                "parser_succeeded": used_parser,
                "parser_error": parser_error,
                "parser_exception_class": parser_exception_class,
                "structural_failure": structural_failure,
                "encrypted": encrypted,
                "signature_hex": signature_hex,
                "byte_count": len(data),
                "meaningful_character_count": meaningful_count,
                "raw_character_count": len(text),
                "normalized_character_count": len(normalized),
                "sufficiency_threshold": MIN_MEANINGFUL_CHARACTERS,
            },
        )

    def _normalize(self, value: str) -> str:
        value = value.replace("\r\n", "\n").replace("\r", "\n")
        value = re.sub(r"[ \t]+", " ", value)
        value = re.sub(r"\n{3,}", "\n\n", value)
        return value.strip()


class DocxExtractionAdapter:
    PARSER_LIBRARY = "python-docx"

    def extract(self, file_obj) -> ExtractionPayload:
        file_obj.seek(0)
        data = file_obj.read()
        signature_hex = data[:4].hex()
        if not data.startswith(b"PK"):
            raise ExtractionError(
                "The uploaded file does not look like a DOCX package.",
                code="invalid_docx_signature",
                details={
                    "adapter": "docx",
                    "parser_library": self.PARSER_LIBRARY,
                    "dependency_available": False,
                    "signature_hex": signature_hex,
                },
            )

        try:
            from docx import Document  # type: ignore
        except Exception as exc:
            raise ExtractionError(
                "The DOCX parser dependency is unavailable.",
                code="parser_dependency_unavailable",
                details={
                    "adapter": "docx",
                    "parser_library": self.PARSER_LIBRARY,
                    "dependency_available": False,
                    "signature_hex": signature_hex,
                    "parser_error": f"{exc.__class__.__name__}: {exc}",
                },
            ) from exc

        if not zipfile.is_zipfile(io.BytesIO(data)):
            raise ExtractionError(
                "The uploaded file is not a valid DOCX container.",
                code="invalid_docx_container",
                details={
                    "adapter": "docx",
                    "parser_library": self.PARSER_LIBRARY,
                    "dependency_available": True,
                    "signature_hex": signature_hex,
                },
            )

        try:
            document = Document(io.BytesIO(data))
        except Exception as exc:
            raise ExtractionError(
                "The DOCX file could not be parsed.",
                code="malformed_docx",
                details={
                    "adapter": "docx",
                    "parser_library": self.PARSER_LIBRARY,
                    "dependency_available": True,
                    "signature_hex": signature_hex,
                    "parser_error": f"{exc.__class__.__name__}: {exc}",
                },
            ) from exc

        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text and paragraph.text.strip()]
        table_rows: list[str] = []
        table_count = 0
        for table in document.tables:
            table_count += 1
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    table_rows.append(" | ".join(cells))

        text_parts = [*paragraphs, *table_rows]
        text = "\n\n".join(text_parts)
        normalized = re.sub(r"\n{3,}", "\n\n", text.replace("\r\n", "\n").replace("\r", "\n")).strip()
        meaningful_count = meaningful_character_count(normalized)
        return ExtractionPayload(
            extracted_text=text,
            normalized_text=normalized,
            extraction_method="docx_text",
            sufficient_text=meaningful_count >= MIN_MEANINGFUL_CHARACTERS,
            metadata={
                "adapter": "docx",
                "parser_library": self.PARSER_LIBRARY,
                "dependency_available": True,
                "signature_hex": signature_hex,
                "paragraph_count": len(paragraphs),
                "table_count": table_count,
                "meaningful_character_count": meaningful_count,
                "raw_character_count": len(text),
                "normalized_character_count": len(normalized),
                "sufficiency_threshold": MIN_MEANINGFUL_CHARACTERS,
            },
        )
